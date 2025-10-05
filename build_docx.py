import json
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import cast

# ---------- Load style spec ----------
STYLE_SPEC = {}

def load_styles(style_path: Path):
    print(f"[DEBUG] Entering load_styles with style_path: {style_path}")
    global STYLE_SPEC
    STYLE_SPEC = json.loads(style_path.read_text(encoding="utf-8"))
    print(f"[DEBUG] STYLE_SPEC loaded with {len(STYLE_SPEC)} styles.")

# ---------- Utilities ----------

# if style does not exist it falls back to the normal style
def style_exists(doc: Document, name: str) -> bool:
    print(f"[DEBUG] Entering style_exists for style: '{name}'")
    try:
        _ = doc.styles[name]
        print(f"[DEBUG] Style '{name}' exists.")
        return True
    except KeyError:
        print(f"[DEBUG] Style '{name}' does NOT exist.")
        return False

def safe_style(doc: Document, name: str, fallback="Normal") -> str:
    print(f"[DEBUG] Entering safe_style for name: '{name}', fallback: '{fallback}'")
    exists = style_exists(doc, name)
    if exists:
        print(f"[DEBUG] safe_style returning original: '{name}'")
        return name
    else:
        fallback_exists = style_exists(doc, fallback)
        if fallback_exists:
            print(f"[DEBUG] safe_style returning fallback: '{fallback}'")
            return fallback
        else:
            default_style = doc.styles[0].name
            print(f"[DEBUG] safe_style returning document default: '{default_style}'")
            return default_style

ALIGN_MAP = {
    "LEFT (0)": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER (1)": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT (2)": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY (3)": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "DISTRIBUTE (4)": WD_ALIGN_PARAGRAPH.DISTRIBUTE if hasattr(WD_ALIGN_PARAGRAPH, "DISTRIBUTE") else WD_ALIGN_PARAGRAPH.JUSTIFY,
    "JUSTIFY_MED (5)": WD_ALIGN_PARAGRAPH.JUSTIFY,  # best-effort
    "JUSTIFY_HI (7)": WD_ALIGN_PARAGRAPH.JUSTIFY,   # best-effort
    "JUSTIFY_LOW (8)": WD_ALIGN_PARAGRAPH.JUSTIFY,  # Word shows "Justify Low"; python-docx uses JUSTIFY
    # tolerate raw short names too:
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# converts name or numbers to actual ALIGN values 
def parse_alignment(val):
    print(f"[DEBUG] Entering parse_alignment with value: {val} (type: {type(val)})")
    if val is None:
        print("[DEBUG] parse_alignment returning None for None input.")
        return None
    if isinstance(val, int):
        # python-docx enums are ints under the hood; accept 0,1,2,3
        try:
            result = WD_ALIGN_PARAGRAPH(val)
            print(f"[DEBUG] parse_alignment returning {result} for integer input.")
            return result
        except Exception as e:
            print(f"[DEBUG] parse_alignment failed for integer {val}: {e}")
            return None
    if isinstance(val, str):
        result = ALIGN_MAP.get(val.strip(), None)
        print(f"[DEBUG] parse_alignment returning {result} for string input '{val}'.")
        return result
    print("[DEBUG] parse_alignment returning None for unhandled type.")
    return None

def set_paragraph_rtl(p, align: WD_ALIGN_PARAGRAPH = None):
    print(f"[DEBUG] Entering set_paragraph_rtl for paragraph with text: {ascii(p.text[:30])}...")
    # force RTL
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        print("[DEBUG] Adding w:bidi element.")
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    print("[DEBUG] Set w:bidi to 1.")
    # alignment
    if align is not None:
        print(f"[DEBUG] Setting alignment to {align}.")
        p.alignment = align
    else:
        # Safe default for Persian body text
        if p.alignment is None:
            print("[DEBUG] Alignment is None, setting to JUSTIFY as default.")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            print(f"[DEBUG] Alignment already set to {p.alignment}, not changing.")

def set_run_lang(run, locale="fa-IR"):
    print(f"[DEBUG] Entering set_run_lang for run with text: {ascii(run.text[:30])}..., locale: {locale}")
    rPr = run._r.get_or_add_rPr()
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        print("[DEBUG] Adding w:lang element.")
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:bidi"), locale)
    lang.set(qn("w:val"), locale)
    print(f"[DEBUG] Set w:lang bidi and val to '{locale}'.")

def add_field_simple(paragraph, instr_text):
    print(f"[DEBUG] Entering add_field_simple with instr_text: '{instr_text}'")
    # PAGE field will inherit paragraph/run formatting and language
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr_text)
    paragraph._p.append(fld)
    print("[DEBUG] Appended w:fldSimple element.")

# ---------- Apply JSON -> document styles (key subset we care about) ----------
# this manages styles 
def set_style_bidi_and_alignment(doc: Document, style_name: str, style_json: dict):
    print(f"[DEBUG] Entering set_style_bidi_and_alignment for style: '{style_name}'")
    if not style_exists(doc, style_name):
        print(f"[DEBUG] Style '{style_name}' not found in document, skipping.")
        return
    st = doc.styles[style_name]
    # paragraph styles only; skip character/table
    if getattr(st, "type", None) and getattr(st.type, "value", None) != 1:
        print(f"[DEBUG] Style '{style_name}' is not a paragraph style (type: {getattr(st, 'type', None)}), skipping.")
        return

    # Paragraph-level alignment from JSON, fallback: RIGHT for headings, JUSTIFY for body
    desired_align = parse_alignment(style_json.get("alignment"))
    if desired_align is None:
        print(f"[DEBUG] No alignment in JSON for '{style_name}'. Falling back.")
        if style_name.startswith("Heading"):
            desired_align = WD_ALIGN_PARAGRAPH.RIGHT
            print("[DEBUG] Fallback for Heading: RIGHT.")
        else:
            # "Normal" etc.
            desired_align = WD_ALIGN_PARAGRAPH.JUSTIFY
            print("[DEBUG] Fallback for Body/Other: JUSTIFY.")
    st.paragraph_format.alignment = desired_align
    print(f"[DEBUG] Set style '{style_name}' alignment to {desired_align}.")

    # Ensure RTL on the style (w:pPr/w:bidi)
    pPr = st.element.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    print(f"[DEBUG] Ensured RTL (w:bidi=1) for style '{style_name}'.")

    # Ensure run language fa-IR on style rPr
    rPr = st.element.get_or_add_rPr()
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:bidi"), "fa-IR")
    lang.set(qn("w:val"), "fa-IR")
    print(f"[DEBUG] Ensured language fa-IR for style '{style_name}'.")

def apply_styles_from_json(doc: Document):
    print("[DEBUG] Entering apply_styles_from_json.")
    # Apply to Normal + Headings + common styles if present in STYLE_SPEC
    targets = set(["Normal", "List Paragraph", "Header", "Footer"])
    targets.update([f"Heading {i}" for i in range(1, 10)])
    print(f"[DEBUG] Target styles to process: {targets}")
    for name in targets:
        spec = STYLE_SPEC.get(name, {})
        print(f"[DEBUG] Processing target style '{name}' with spec from JSON: {spec != {}}")
        set_style_bidi_and_alignment(doc, name, spec)
    print("[DEBUG] Finished apply_styles_from_json.")

# ---------- Numbering: force RTL & fa-IR at levels ----------
def ensure_numbering_rtl(doc: Document):
    print("[DEBUG] Entering ensure_numbering_rtl.")
    numpart = getattr(doc.part, "numbering_part", None)
    if numpart is None:
        print("[DEBUG] No numbering part found in document. Skipping.")
        return
    root = numpart.element
    # For each level, set pPr/jc=right and rPr/lang bidi=fa-IR
    levels_found = root.findall(".//w:lvl", namespaces=root.nsmap)
    print(f"[DEBUG] Found {len(levels_found)} numbering levels to process.")
    for i, lvl in enumerate(levels_found):
        print(f"[DEBUG] Processing level {i+1}...")
        # pPr/jc
        pPr = lvl.find("./w:pPr", namespaces=root.nsmap)
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            lvl.append(pPr)
        jc = pPr.find("./w:jc", namespaces=root.nsmap)
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "right")
        print("[DEBUG] Set numbering level alignment to 'right'.")

        # rPr/lang
        rPr = lvl.find("./w:rPr", namespaces=root.nsmap)
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            lvl.append(rPr)
        lang = rPr.find("./w:lang", namespaces=root.nsmap)
        if lang is None:
            lang = OxmlElement("w:lang")
            rPr.append(lang)
        lang.set(qn("w:bidi"), "fa-IR")
        lang.set(qn("w:val"), "fa-IR")
        print("[DEBUG] Set numbering level language to 'fa-IR'.")
    print("[DEBUG] Finished ensure_numbering_rtl.")

# ---------- Content helpers (use document styles, not JSON names) ----------

def add_para(doc, text, style_name="Normal"):
    print(f"[DEBUG] Entering add_para with style '{style_name}', text: {ascii(text[:50])}...")
    final_style_name = safe_style(doc, style_name)
    print(f"[DEBUG] Resolved style for add_para: '{final_style_name}'")
    p = doc.add_paragraph(text, style=final_style_name)
    # Use the style's alignment unless JSON dictates otherwise
    json_style = STYLE_SPEC.get(final_style_name, {})
    align = parse_alignment(json_style.get("alignment"))
    set_paragraph_rtl(p, align)
    for r in p.runs:
        set_run_lang(r)
    return p

def add_heading(doc, text, level=1):
    print(f"[DEBUG] Entering add_heading with level {level}, text: {ascii(text[:50])}...")
    style_name = f"Heading {level}"
    final_style_name = safe_style(doc, style_name)
    print(f"[DEBUG] Resolved style for add_heading: '{final_style_name}'")
    p = doc.add_paragraph(text, style=final_style_name)
    # Headings default to RIGHT unless JSON says otherwise
    json_style = STYLE_SPEC.get(final_style_name, {})
    align = parse_alignment(json_style.get("alignment")) or WD_ALIGN_PARAGRAPH.RIGHT
    print(f"[DEBUG] Alignment for heading: {align}")
    set_paragraph_rtl(p, align)
    for r in p.runs:
        set_run_lang(r)
    return p

def add_list_item(doc, text, list_type="ol", meta=None):
    print(f"[DEBUG] Entering add_list_item for type '{list_type}', text: {ascii(text[:50])}...")
    style_map = (meta or {}).get("listStyleMap", {})
    desired = style_map.get(list_type)
    print(f"[DEBUG] Desired list style from map: '{desired}'")

    # fallbacks if style not mapped or missing in template
    if not desired:
        desired = "List Paragraph" if list_type == "ol" else "Bulet"
        print(f"[DEBUG] No desired style, falling back to '{desired}'")

    style_name = desired if style_exists(doc, desired) else (
        "List Paragraph" if style_exists(doc, "List Paragraph") else safe_style(doc, "Normal")
    )
    print(f"[DEBUG] Final resolved list style: '{style_name}'")

    p = doc.add_paragraph(text, style=style_name)

    # alignment from JSON (if any), else leave style's alignment
    json_style = STYLE_SPEC.get(style_name, {}) if "STYLE_SPEC" in globals() else {}
    align = parse_alignment(json_style.get("alignment"))
    set_paragraph_rtl(p, align)

    for r in p.runs:
        set_run_lang(r)
    return p
def clear_paragraphs(container):
    print(f"[DEBUG] Entering clear_paragraphs for a container (e.g., header/footer).")
    # Remove all existing paragraphs from header/footer to avoid mixed formatting
    num_paras = len(list(container.paragraphs))
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)
    print(f"[DEBUG] Removed {num_paras} paragraphs.")

def apply_header_footer(doc, section, cfg):
    print("[DEBUG] Entering apply_header_footer.")
    # header
    if cfg.get("enabled"):
        print("[DEBUG] Header is enabled.")
        hdr = section.header
        clear_paragraphs(hdr)
        desired_align = parse_alignment(cfg.get("align")) or WD_ALIGN_PARAGRAPH.RIGHT
        print(f"[DEBUG] Header alignment: {desired_align}")
        for i, rn in enumerate(cfg.get("runs", [])):
            print(f"[DEBUG] Processing header run {i+1}...")
            p = hdr.add_paragraph(style=safe_style(doc, rn.get("style", "Header")))
            set_paragraph_rtl(p, desired_align)
            if "text" in rn:
                print(f"[DEBUG] Adding header text: {ascii(rn['text'])}")
                set_run_lang(p.add_run(rn["text"]))
            if rn.get("field") == "pageNumber":
                print("[DEBUG] Adding page number field to header.")
                add_field_simple(p, "PAGE")

    # footer
    if cfg.get("enabled"):
        print("[DEBUG] Footer is enabled.")
        ftr = section.footer
        clear_paragraphs(ftr)
        desired_align = parse_alignment(cfg.get("align")) or WD_ALIGN_PARAGRAPH.RIGHT
        print(f"[DEBUG] Footer alignment: {desired_align}")
        for i, rn in enumerate(cfg.get("runs", [])):
            print(f"[DEBUG] Processing footer run {i+1}...")
            p = ftr.add_paragraph(style=safe_style(doc, rn.get("style", "Footer")))
            set_paragraph_rtl(p, desired_align)
            if "text" in rn:
                print(f"[DEBUG] Adding footer text: {ascii(rn['text'])}")
                set_run_lang(p.add_run(rn["text"]))
            if rn.get("field") == "pageNumber":
                print("[DEBUG] Adding page number field to footer.")
                add_field_simple(p, "PAGE")

def write_section(doc, meta, node):
    print("[DEBUG] Entering write_section.")
    sect_cfg = node.get("section") or {}
    break_kind = sect_cfg.get("break", "oddPage")
    start_type = WD_SECTION.ODD_PAGE if break_kind == "oddPage" else WD_SECTION.NEW_PAGE
    print(f"[DEBUG] Adding new section with break type: {break_kind} ({start_type})")
    section = doc.add_section(start_type=start_type)

    # Note: The original code had a bug applying header config to footer. Correcting it.
    header_cfg = node.get("header", {})
    if header_cfg:
        print("[DEBUG] Applying header configuration.")
        apply_header_footer_specific(doc, section, header_cfg, is_header=True)

    footer_cfg = node.get("footer", {})
    if footer_cfg:
        print("[DEBUG] Applying footer configuration.")
        apply_header_footer_specific(doc, section, footer_cfg, is_header=False)


    ch = node.get("chapter")
    if ch:
        print(f"[DEBUG] Writing chapter: {ascii(ch.get('title', 'Untitled'))}")
        add_heading(doc, ch["title"], level=1)
        for i, par in enumerate(ch.get("intro", [])):
            print(f"[DEBUG] Writing intro paragraph {i+1}.")
            # allow both {"text": "..."} or plain strings in your content JSON
            text = par.get("text") if isinstance(par, dict) else str(par)
            add_para(doc, text, style_name=meta.get("defaultParagraphStyle", "Normal"))
        for i, sec in enumerate(ch.get("sections", [])):
            print(f"[DEBUG] Writing chapter subsection {i+1}.")
            write_subsection(doc, meta, sec)
    print("[DEBUG] Finished write_section.")

def apply_header_footer_specific(doc, section, cfg, is_header):
    """A corrected helper to apply config to either header or footer."""
    print(f"[DEBUG] Entering apply_header_footer_specific for {'header' if is_header else 'footer'}.")
    if cfg.get("enabled"):
        container = section.header if is_header else section.footer
        default_style = "Header" if is_header else "Footer"
        print(f"[DEBUG] {'Header' if is_header else 'Footer'} is enabled.")
        clear_paragraphs(container)
        desired_align = parse_alignment(cfg.get("align")) or WD_ALIGN_PARAGRAPH.RIGHT
        print(f"[DEBUG] Alignment: {desired_align}")
        for i, rn in enumerate(cfg.get("runs", [])):
            print(f"[DEBUG] Processing run {i+1}...")
            p = container.add_paragraph(style=safe_style(doc, rn.get("style", default_style)))
            set_paragraph_rtl(p, desired_align)
            if "text" in rn:
                print(f"[DEBUG] Adding text: {ascii(rn['text'])}")
                set_run_lang(p.add_run(rn["text"]))
            if rn.get("field") == "pageNumber":
                print("[DEBUG] Adding page number field.")
                add_field_simple(p, "PAGE")

def write_subsection(doc, meta, sec):
    level = sec.get("level", 2)
    print(f"[DEBUG] Entering write_subsection for level {level}, title: {ascii(sec.get('title', 'Untitled'))}")
    add_heading(doc, sec["title"], level=level)
    for i, node in enumerate(sec.get("content", [])):
        print(f"[DEBUG] Processing content node {i+1} in subsection.")
        if isinstance(node, dict) and "text" in node:
            print("[DEBUG] Content node is a paragraph.")
            add_para(doc, node["text"], style_name=meta.get("defaultParagraphStyle", "Normal"))
        elif isinstance(node, dict) and "list" in node:
            print("[DEBUG] Content node is a list.")
            for j, item in enumerate(node["list"].get("items", [])):
                print(f"[DEBUG] Adding list item {j+1}.")
                itxt = item.get("text") if isinstance(item, dict) else str(item)
                add_list_item(doc, itxt, list_type=node["list"].get("type", "ol"), meta=meta)
        elif isinstance(node, str):
            print("[DEBUG] Content node is a plain string paragraph.")
            add_para(doc, node, style_name=meta.get("defaultParagraphStyle", "Normal"))
    for i, sub in enumerate(sec.get("sections", [])):
        print(f"[DEBUG] Writing nested subsection {i+1}.")
        write_subsection(doc, meta, sub)
    print("[DEBUG] Finished write_subsection.")

# ---------- Main ----------

if __name__ == "__main__":
    print("[DEBUG] Script started.")
    if len(sys.argv) < 4:
        print("Usage: python build_docx.py <styles.json> <content.json> <out.docx>")
        sys.exit(1)

    style_path = Path(sys.argv[1])
    content_path = Path(sys.argv[2])
    out_path = Path(sys.argv[3])
    print(f"[DEBUG] Style path: {style_path}")
    print(f"[DEBUG] Content path: {content_path}")
    print(f"[DEBUG] Output path: {out_path}")

    load_styles(style_path)
    print("[DEBUG] Loading content JSON...")
    spec = json.loads(content_path.read_text(encoding="utf-8"))
    print("[DEBUG] Content JSON loaded.")

    # Open a seed docx if provided (recommended: a .docx saved from your .dotx)
    # tmpl = spec.get("meta", {}).get("template")
    # if tmpl and Path(tmpl).is_file():
    #     print(f"[DEBUG] Loading the Template Document from: {tmpl}")
    #     doc = Document(tmpl)
    # else:
    print("[DEBUG] Creating new blank Document.")
    doc = Document()

    # Apply JSON style intentions to the actual document styles
    print("[DEBUG] Starting to apply styles from JSON to document.")
    apply_styles_from_json(doc)

    # Make multilevel numbering RTL + fa-IR at all levels
    print("[DEBUG] Starting to ensure numbering is RTL.")
    ensure_numbering_rtl(doc)

    # Build content
    print("[DEBUG] Starting to build content from spec.")
    # The root of the spec is treated as the first section's content
    write_section(doc, spec.get("meta", {}), spec)

    print(f"[DEBUG] Saving document to {out_path}...")
    doc.save(out_path)
    print(f"Wrote {out_path}")
    print("[DEBUG] Script finished.")
