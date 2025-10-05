
# def add_toc(paragraph):
#     """Insert a TOC field that Word will auto-generate."""
#     run = paragraph.add_run()

#     fld_char_begin = OxmlElement("w:fldChar")
#     fld_char_begin.set(qn("w:fldCharType"), "begin")

#     instr_text = OxmlElement("w:instrText")
#     instr_text.set(qn("xml:space"), "preserve")
#     # \o "1-3" → include Heading 1–3
#     # \h → hyperlinks
#     # \z → hide page numbers in web layout
#     # \u → use outline levels
#     instr_text.text = r'TOC \o "1-3" \h \z \u'

#     fld_char_sep = OxmlElement("w:fldChar")
#     fld_char_sep.set(qn("w:fldCharType"), "separate")

#     fld_char_end = OxmlElement("w:fldChar")
#     fld_char_end.set(qn("w:fldCharType"), "end")

#     run._r.extend([fld_char_begin, instr_text, fld_char_sep, fld_char_end])

from docx import Document

doc = Document("out-template.docx")



toc1 = doc.styles["toc 1"]
toc2 = doc.styles["toc 2"]
toc3 = doc.styles["toc 3"]

# Title for TOC
doc.add_paragraph("فهرست مطالب", style=toc1)

# Add TOC placeholder


# Add some sample headings
doc.add_paragraph("فصل چهارم: سبک ها و قلم ها", style=toc1)
doc.add_paragraph("قلم‌های فارسی", style=toc2)
doc.add_paragraph("قلم‌های انگلیسی", style=toc2)
doc.add_paragraph("فاصله‌ها (روابط ریاضی)", style=toc2)
doc.add_paragraph("فاصله‌های افقی و عمودی", style=toc2)
doc.add_paragraph("فاصله کلی از چهار طرف کاغذ", style=toc3)

doc.save("with_toc.docx")