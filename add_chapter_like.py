# add_chapter_like.py
# pip install python-docx

from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
import json
import re
import unicodedata
import argparse

# ----------------- Model -----------------

@dataclass
class Node:
    title: str
    level: int
    start_idx: int
    content_idxs: List[int] = field(default_factory=list)
    children: List["Node"] = field(default_factory=list)

# ----------------- Detectors & Normalizers -----------------

def _heading_level(p: Paragraph) -> Optional[int]:
    name = (p.style.name or "").lower()
    m = re.search(r'(heading|überschrift)\s*(\d+)$', name)
    if m:
        return int(m.group(2))
    pPr = p._p.pPr
    if pPr is not None and pPr.outlineLvl is not None:
        return int(pPr.outlineLvl.val) + 1
    return None

def has_sectPr(paragraph: Paragraph) -> bool:
    pPr = paragraph._p.pPr
    return (pPr is not None) and (pPr.sectPr is not None)

FM_TITLES = {"فهرست مطالب","فهرست جداول","فهرست اشکال","فهرست شکل‌ها",
             "Table of Contents","List of Tables","List of Figures"}

def normalize(s: str) -> str:
    if s is None:
        return ""
    s = s.replace("\\n", " ")  # literal \n -> space
    s = unicodedata.normalize("NFKC", s)
    s = s.translate(str.maketrans({"ي": "ی", "ك": "ک", "ۀ": "ه", "أ": "ا", "إ": "ا", "ٱ": "ا"}))
    s = s.replace("\u200c", "").replace("\u200f", "").replace("\u200e", "")
    s = re.sub(r"\s+", " ", s.strip())
    return s.casefold()

def is_front_matter_title(p: Paragraph) -> bool:
    t = normalize(p.text)
    if t in {x.casefold() for x in FM_TITLES}:
        return True
    s = (p.style.name or "").lower()
    if any(k in s for k in ["toc heading","table of contents","list of tables","list of figures"]):
        return True
    return False

# ----------------- Tree (section-aware & TOC/LOF/LOT-aware) -----------------

def build_tree(doc: Document) -> List[Node]:
    roots: List[Node] = []
    stack: List[Node] = []

    def start_root_node(title: str, idx: int):
        n = Node(title=title, level=1, start_idx=idx)
        roots.append(n)
        stack.clear()
        stack.append(n)

    for idx, p in enumerate(doc.paragraphs):
        # Treat TOC/LOF/LOT titles as their own top-level nodes
        if is_front_matter_title(p):
            start_root_node(p.text.strip(), idx)
            continue

        lvl = _heading_level(p)
        if lvl is not None:
            node = Node(title=p.text.strip(), level=lvl, start_idx=idx)
            while stack and stack[-1].level >= lvl:
                stack.pop()
            if stack:
                stack[-1].children.append(node)
            else:
                roots.append(node)
            stack.append(node)
        else:
            if stack:
                stack[-1].content_idxs.append(idx)
                # stop capturing at section end to avoid spillover
                if has_sectPr(p):
                    stack.clear()
    return roots

def iter_nodes(nodes: List[Node]) -> List[Node]:
    out = []
    def walk(n: Node):
        out.append(n)
        for c in n.children:
            walk(c)
    for r in nodes:
        walk(r)
    return out

def find_chapter(tree: List[Node], title: str) -> Optional[Node]:
    want = normalize(title)
    cands = [n for n in iter_nodes(tree) if n.level == 1]
    exact = [n for n in cands if normalize(n.title) == want]
    if exact:
        return exact[0]
    partial = [n for n in cands if want and want in normalize(n.title)]
    if len(partial) == 1:
        return partial[0]
    return None

# ----------------- Formatting helpers (for add mode) -----------------

def add_paragraph_after(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_para.style = style  # style object or name
        except Exception:
            pass
    if text:
        new_para.add_run(text)  # no direct run formatting
    return new_para

def add_body_block(doc: Document, text: str, style=None):
    blocks = [t.strip() for t in re.split(r"\n\s*\n", text or "") if t.strip()]
    for blk in blocks:
        p = doc.add_paragraph()
        if style is not None:
            try:
                p.style = style
            except Exception:
                pass
        p.add_run(blk)

def infer_node_body_style(doc: Document, node: Node):
    for i in node.content_idxs:
        if 0 <= i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            if _heading_level(p) is None and p.text.strip():
                try:
                    return p.style
                except Exception:
                    pass
    try:
        return doc.styles["Normal"]
    except Exception:
        return None

# ----------------- JSON skeleton (for export mode) -----------------

def collapse_content(doc: Document, par_idxs: List[int]) -> str:
    """Join the node's body paragraphs into a single string with blank-line separators."""
    parts = []
    for i in par_idxs:
        if 0 <= i < len(doc.paragraphs):
            txt = doc.paragraphs[i].text
            if txt.strip():
                parts.append(txt)
    return "\n\n".join(parts)

def node_to_json(doc: Document, node: Node, blank_content: bool = False) -> Dict[str, Any]:
    d: Dict[str, Any] = {"__content__": "" if blank_content else collapse_content(doc, node.content_idxs)}
    for c in node.children:
        d[c.title] = node_to_json(doc, c, blank_content=blank_content)
    return d

def export_chapter_skeleton(doc: Document, chapter_node: Node, output_json_path: str, blank_content: bool):
    data = {chapter_node.title: node_to_json(doc, chapter_node, blank_content=blank_content)}
    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ----------------- Add a new chapter (existing behavior) -----------------

def create_chapter_from_template(
    doc: Document,
    template_node: Node,
    new_title: str,
    content_spec: Dict[str, Any] = None,
    insert_mode: str = "pagebreak"
):
    content_spec = content_spec or {}

    # 1) Page/Section handling (keep headers/footers identical)
    if insert_mode == "section":
        sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
        for part in ("header","footer","first_page_header","even_page_header","first_page_footer","even_page_footer"):
            try:
                getattr(sec, part).is_linked_to_previous = True
            except Exception:
                pass
    else:
        doc.add_page_break()

    # 2) Chapter heading uses EXACT style of the template chapter heading
    template_heading_style = doc.paragraphs[template_node.start_idx].style
    h1 = doc.add_paragraph()
    try:
        h1.style = template_heading_style
    except Exception:
        pass
    h1.add_run(new_title)

    # 3) Body style under chapter inherits from template chapter's body
    chapter_body_style = infer_node_body_style(doc, template_node)
    if isinstance(content_spec.get("__content__"), str):
        add_body_block(doc, content_spec["__content__"], style=chapter_body_style)

    # 4) Recursively add child headings with their exact styles
    def add_children(template_children: List[Node], spec: Dict[str, Any]):
        provided: Dict[str, Any] = {}
        for k, v in spec.items():
            if k == "__content__":
                continue
            provided[normalize(k)] = {"__content__": v} if isinstance(v, str) else v

        for child in template_children:
            c_spec = provided.get(normalize(child.title), {})
            child_heading_style = doc.paragraphs[child.start_idx].style
            new_child_title = c_spec.get("__title__", child.title)

            h = doc.add_paragraph()
            try:
                h.style = child_heading_style
            except Exception:
                pass
            h.add_run(new_child_title)

            child_body_style = infer_node_body_style(doc, child)
            if isinstance(c_spec.get("__content__"), str):
                add_body_block(doc, c_spec["__content__"], style=child_body_style)

            if child.children:
                add_children(child.children, c_spec if isinstance(c_spec, dict) else {})

    add_children(template_node.children, content_spec)

# ----------------- CLI -----------------

def main():
    ap = argparse.ArgumentParser(
        description="Add a new chapter from an existing chapter's structure OR export a chapter's skeleton to JSON (section-aware; preserves headers/footers; inherits styles)."
    )
    ap.add_argument("input_docx", help="Input .docx")
    ap.add_argument("template_chapter_title", nargs="?", default=None,
                    help='Existing chapter title (H1). Accepts literal \\n or a unique substring. If using --template-index, this can be omitted.')
    # These become optional when you only export:
    ap.add_argument("new_chapter_title", nargs="?", default=None, help="New chapter title (H1) for add mode")
    ap.add_argument("output_docx", nargs="?", default=None, help="Output .docx for add mode")

    # Modes
    ap.add_argument("--export-skeleton", metavar="JSON_PATH",
                    help="Export the selected chapter's skeleton to this JSON file and exit (no add).")
    ap.add_argument("--blank-content", action="store_true",
                    help="When exporting skeleton, set all __content__ fields to empty strings.")
    ap.add_argument("--json", help="JSON with body text for the NEW chapter & sub-sections (add mode)", default=None)
    ap.add_argument("--insert-mode", choices=["pagebreak","section"], default="pagebreak",
                    help="Add mode: pagebreak = same section; section = new section but headers/footers linked")

    # Selection helpers
    ap.add_argument("--list-chapters", action="store_true", help="List H1 titles with indices and exit")
    ap.add_argument("--template-index", type=int, default=None,
                    help="Pick the template chapter by index from --list-chapters")

    args = ap.parse_args()

    doc = Document(args.input_docx)
    tree = build_tree(doc)
    h1_nodes = [n for n in tree if n.level == 1]

    if args.list_chapters:
        for i, n in enumerate(h1_nodes):
            print(f"[{i}] {n.title}")
        return

    # Pick the chapter node
    if args.template_index is not None:
        if 0 <= args.template_index < len(h1_nodes):
            template = h1_nodes[args.template_index]
        else:
            raise SystemExit(f"Invalid --template-index (0..{len(h1_nodes)-1})")
    else:
        if not args.template_chapter_title:
            raise SystemExit("Please supply template_chapter_title or use --template-index.")
        template = find_chapter(tree, args.template_chapter_title)

    if not template:
        raise SystemExit(f"Template chapter not found: {args.template_chapter_title}")

    # ------- Export skeleton mode -------
    if args.export_skeleton:
        export_chapter_skeleton(doc, template, args.export_skeleton, blank_content=args.blank_content)
        print(f"Wrote skeleton for '{template.title}' -> {args.export_skeleton}")
        return

    # ------- Add chapter mode -------
    if not args.new_chapter_title or not args.output_docx:
        raise SystemExit("Add mode requires: new_chapter_title and output_docx (or use --export-skeleton to export only).")

    content_spec = {}
    if args.json:
        with open(args.json, "r", encoding="utf-8") as f:
            content_spec = json.load(f)

    create_chapter_from_template(
        doc,
        template_node=template,
        new_title=args.new_chapter_title,
        content_spec=content_spec,
        insert_mode=args.insert_mode,
    )
    doc.save(args.output_docx)
    print(f"Added '{args.new_chapter_title}' based on '{template.title}'. Saved: {args.output_docx}")

if __name__ == "__main__":
    main()
