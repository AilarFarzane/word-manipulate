# add_custom_chapter_strict.py
# pip install python-docx

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import argparse, json, re

def add_field(paragraph, instr: str):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = ""
    r.append(t); fld.append(r)
    paragraph._p.append(fld)

def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)

def set_paragraph_rtl(p, rtl=True):
    """Force a paragraph to Right-to-Left (w:bidi=1)."""
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if rtl:
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
        bidi.set(qn('w:val'), '1')
    else:
        if bidi is not None:
            pPr.remove(bidi)

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if text:
        run = p.add_run(text)
        if p.style and p.style.font:
            # Attempt to inherit font from style
            run.font.name = p.style.font.name
            # For complex scripts like Arabic/Persian, also set this
            run.font.rtl = True
    set_paragraph_rtl(p, True)
    try:
        sname = (p.style.name or "").lower()
    except Exception:
        sname = ""
    if re.search(r'(heading|überschrift|عنوان)\s*\d*$', sname):  # English, German, Persian UIs
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p



def _heading_level(p):
    name = (p.style.name or "").lower()
    m = re.search(r'(heading|überschrift|عنوان)\s*(\d+)$', name)
    if m: return int(m.group(2))
    return None

def build_tree(doc):
    roots, stack = [], []
    for idx, p in enumerate(doc.paragraphs):
        lvl = _heading_level(p)
        if lvl:
            n = {"title": p.text.strip(), "level": lvl, "idx": idx, "children": []}
            while stack and stack[-1]["level"] >= lvl: stack.pop()
            if stack: stack[-1]["children"].append(n)
            else: roots.append(n)
            stack.append(n)
    return roots

def find_chapter_by_title(doc, title: str):
    tree = build_tree(doc)
    want = re.sub(r"\s+", " ", (title or "").strip()).casefold()
    for n in tree:
        if n["level"] == 1 and re.sub(r"\s+", " ", n["title"]).casefold() == want:
            return n
    return None

def get_heading_styles(doc, template_node):
    """
    Infers heading styles. If a template_node (from a template chapter)
    is provided, it traverses that chapter to find the actual styles used
    for each heading level. It falls back to default names like "Heading 1"
    for levels not found in the template.
    """
    hstyles = {}
    if not template_node:
        # Fallback to default names if no template is provided
        for lvl in range(1, 7):
            try:
                hstyles[lvl] = doc.styles[f"Heading {lvl}"]
                continue
            except KeyError:
                pass
            # If not found, search all styles for a match
            for style in doc.styles:
                m = re.search(r'(heading|überschrift|عنوان)\s*(' + str(lvl) + r')$', (style.name or "").lower())
                if m:
                    hstyles[lvl] = style
                    break
        return hstyles

    # BFS to find the first instance of a style for each heading level
    queue = [template_node]
    visited_levels = set()
    while queue:
        node = queue.pop(0)
        lvl = node.get("level")
        if lvl and lvl not in visited_levels:
            try:
                hstyles[lvl] = doc.paragraphs[node["idx"]].style
                visited_levels.add(lvl)
            except IndexError: pass
        
        if "children" in node:
            queue.extend(node["children"])

    # For any levels not found in the template, fall back to default names
    for lvl in range(1, 7):
        if lvl not in hstyles:
            try: hstyles[lvl] = doc.styles[f"Heading {lvl}"]
            except KeyError: pass
    return hstyles


def infer_body_style(doc):
    try: return doc.styles["Normal"]
    except: return None

def set_header_footer(sec, header_data, footer_data):
    # Header
    if not header_data.get("link_to_previous", False):
        try: sec.header.is_linked_to_previous = False
        except: pass
        header_text = header_data.get("text") if header_data.get("enabled", True) else None
        if header_text:
            h = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
            clear_runs(h); h.add_run(header_text)
            header_align = (header_data.get("align") or "center").lower()
            h.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                           "center": WD_ALIGN_PARAGRAPH.CENTER,
                           "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(header_align, WD_ALIGN_PARAGRAPH.CENTER)
            set_paragraph_rtl(h)

    # Footer
    if not footer_data.get("link_to_previous", False):
        try: sec.footer.is_linked_to_previous = False
        except: pass
        footer_note = footer_data.get("note")
        show_page_number = bool(footer_data.get("include_page_number", True))
        if footer_note is not None or show_page_number:
            f = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
            clear_runs(f)
            if footer_note:
                f.add_run(footer_note)
                if show_page_number: f.add_run("  ")
            if show_page_number: add_field(f, " PAGE ")
            page_align = (footer_data.get("page_number_align") or "center").lower()
            f.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                           "center": WD_ALIGN_PARAGRAPH.CENTER,
                           "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(page_align, WD_ALIGN_PARAGRAPH.CENTER)
            set_paragraph_rtl(f)

def ensure_paras(x):
    if x is None: return []
    if isinstance(x, str): return [x.strip()]
    if isinstance(x, list): return [str(t).strip() for t in x if str(t).strip()]
    return []

def add_sections(doc, sections, hstyles, body_style):
    for s in sections:
        lvl = int(s.get("level", 2))
        title = s.get("title", "")
        content = ensure_paras(s.get("content"))
        add_paragraph(doc, title, style=hstyles.get(lvl))
        for para in content: add_paragraph(doc, para, style=body_style)
        add_sections(doc, s.get("sections", []), hstyles, body_style)

def build_from_json(doc, data, template_doc, template_node=None, debug=False):
    meta, header, footer, chapter = data["meta"], data["header"], data["footer"], data["chapter"]

    if meta.get("insert_mode","section") == "section":
        sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
        set_header_footer(sec, header, footer)
    else:
        doc.add_page_break()

    hstyles = get_heading_styles(template_doc, template_node)
    body_style = infer_body_style(doc)

    # H1
    add_paragraph(doc, chapter["title"], style=hstyles[1])
    if debug: print(f"[ADD] H1: {chapter['title']}")

    for para in ensure_paras(chapter.get("intro")):
        add_paragraph(doc, para, style=body_style)

    add_sections(doc, chapter.get("sections", []), hstyles, body_style)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input_docx")
    ap.add_argument("chapter_json")
    ap.add_argument("output_docx")
    ap.add_argument("--inherit-from-title", default=None)
    ap.add_argument("--debug", action="store_true")
    args = ap.parse_args()

    doc = Document(args.input_docx)
    with open(args.chapter_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    template_doc = doc
    template_node = None
    
    # Check for an external template document in the JSON
    if "template_document" in data.get("meta", {}):
        template_path = data["meta"]["template_document"]
        try:
            template_doc = Document(template_path)
            if args.inherit_from_title:
                template_node = find_chapter_by_title(template_doc, args.inherit_from_title)
        except Exception as e:
            print(f"Warning: Could not load template '{template_path}'. Using base document for styles. Error: {e}")
    elif args.inherit_from_title:
        template_node = find_chapter_by_title(doc, args.inherit_from_title)

    build_from_json(doc, data, template_doc, template_node, debug=args.debug)
    doc.save(args.output_docx)
    print(f"Saved: {args.output_docx}")

if __name__ == "__main__":
    main()
