from docx import Document
import json
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


with open("content/01-chapter-01.json", "r", encoding="utf-8") as f:
      data = json.load(f)
with open("content/00-frontmatter.json", "r", encoding="utf-8") as f:
    front_json = json.load(f)

doc = Document("template/blank-template.docx")

#helper
def add_bottom_border(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')   # line style
    bottom.set(qn('w:sz'), '9')        # line thickness
    bottom.set(qn('w:space'), '1')      # spacing
    bottom.set(qn('w:color'), '000000') # black
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_page_number(paragraph):
    """Insert a PAGE field into the given paragraph."""
    run = paragraph.add_run()
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

def set_page_number_format(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)

    pgNumType.set(qn("w:fmt"), fmt)

    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    else:
        # remove w:start so numbering continues automatically
        if pgNumType.get(qn("w:start")):
            pgNumType.attrib.pop(qn("w:start"))



# front page replacement
front = front_json["frontMatter"]
for p in doc.paragraphs:
        for key, value in front.items():
            placeholder = f"{{{{{key}}}}}"   # makes {{key}}
            if placeholder in p.text:
                # replace text for the whole paragraph
                p.text = p.text.replace(placeholder, value)


# section 1
section_1  = doc.add_section()
section_1.header.is_linked_to_previous = False
section_1.footer.is_linked_to_previous = False


# remove borders for this section
sectPr = section_1._sectPr
pgBorders = sectPr.find(qn('w:pgBorders'))
if pgBorders is not None:
    sectPr.remove(pgBorders)

# header
header_1_text = section_1.header.add_paragraph(
      data["header"]["text"],
      style=doc.styles[data["header"]["style"]]
)
add_bottom_border(header_1_text)
header_1_text.paragraph_format.space_after = Pt(12)  #one line break at the bottom of the header 

# footer
footer_section_1 = section_1.footer
p = footer_section_1.paragraphs[0]
p.add_run()
add_page_number(p)
set_page_number_format(section_1, fmt="decimal", start=1)


# intro
doc.add_paragraph(
      data["chapter"]["title"],
      style=doc.styles[data["chapter"]["style"]]
)
for para in data["chapter"]["intro"]:
    doc.add_paragraph(
        para["text"],
        style=doc.styles[para["style"]]
    )

#sections for chapter 1
for section in data["chapter"]["sections"]:
      doc.add_paragraph(
            section["title"],
            style=doc.styles[section["style"]]
      )
      #contents of section
          # Section content
      for block in section.get("content", []):
        # Normal paragraph
        if "text" in block:
            doc.add_paragraph(
                block["text"],
                style=doc.styles[block.get("style")]
            )

        # List of items
        if "list" in block:
            for item in block["list"]:
                p = doc.add_paragraph(
                    item["text"],
                    style=doc.styles[item.get("style")]
                )
        if "upload" in block:
            for upload_item in block["upload"]:
                if "image" in upload_item:
                    p = doc.add_paragraph(style=upload_item.get("style"))
                    run = p.add_run()
                    run.add_picture(upload_item["image"])

                if "text" in upload_item:
                    p = doc.add_paragraph(upload_item["text"], style=doc.styles[upload_item.get("style")])

      # --- SUB-SECTIONS ---
      for sub in section.get("sub_sections", []):
        # Subsection title
        doc.add_paragraph(
            sub["title"],
            style=doc.styles[sub.get("style")]
        )
        # Subsection content
        for block in sub.get("content", []):
            if "text" in block:
                  doc.add_paragraph(
                  block["text"],
                  style=doc.styles[block.get("style")]
                  )

        # List of items
            elif "list" in block:
                  for item in block["list"]:
                        p = doc.add_paragraph(
                              item["text"],
                              style=doc.styles[item.get("style")]
                        )

# section 2
section_2 = doc.add_section()

section_2.header.is_linked_to_previous = False
section_2.footer.is_linked_to_previous = False

header_2_text = section_2.header.add_paragraph(
      text="1",
      style=doc.styles[data["header"]["style"]]
)
add_bottom_border(header_2_text)
header_2_text.paragraph_format.space_after = Pt(12)

# footer
footer_section_2 = section_2.footer
p = footer_section_2.paragraphs[0]
p.add_run()
add_page_number(p)
set_page_number_format(section_2, fmt="decimal")
     

doc.save("out-template.docx")
