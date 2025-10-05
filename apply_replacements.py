# apply_replacements.py
# pip install python-docx

from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json, re, unicodedata

@dataclass
class Node:
    title: str
    level: int
    start_idx: int
    content_idxs: List[int] = field(default_factory=list)
    children: List["Node"] = field(default_factory=list)

# ----------------- detectors -----------------
def _heading_level(p: Paragraph) -> Optional[int]:
    name = (p.style.name or "").lower()
    m = re.search(r'(heading|überschrift)\s*(\d+)$', name)
    if m: return int(m.group(2))
    pPr = p._p.pPr
    if pPr is not None and pPr.outlineLvl is not None:
        return int(pPr.outlineLvl.val) + 1
    return None

def has_sectPr(p: Paragraph) -> bool:
    pPr = p._p.pPr
    return (pPr is not None) and (pPr.sectPr is not None)

FM_TITLES = {"فهرست مطالب","فهرست جداول","فهرست اشکال","فهرست شکل‌ها",
             "Table of Contents","List of Tables","List of Figures"}

def normalize_text(s: str) -> str:
    if s is None: return ""
    s = s.replace("\\n"," ")
    s = unicodedata.normalize("NFKC", s)
    s = s.translate(str.maketrans({"ي":"ی","ك":"ک","ۀ":"ه","أ":"ا","إ":"ا","ٱ":"ا"}))
    s = s.replace("\u200c","").replace("\u200f","").replace("\u200e","")
    return re.sub(r"\s+"," ",s.strip()).casefold()

def is_front_matter_title(p: Paragraph) -> bool:
    t = normalize_text(p.text)
    if t in {x.casefold() for x in FM_TITLES}: return True
    s = (p.style.name or "").lower()
    if any(k in s for k in ["toc heading","table of contents","list of tables","list of figures"]):
        return True
    return False

# ----------------- tree (section-aware & TOC/LOF/LOT-aware) -----------------
def build_tree(doc: Document) -> List[Node]:
    roots: List[Node] = []
    stack: List[Node] = []

    def start_root_node(title: str, idx: int):
        n = Node(title=title, level=1, start_idx=idx)
        roots.append(n)
        stack.clear(); stack.append(n)

    for idx, p in enumerate(doc.paragraphs):
        if is_front_matter_title(p):
            start_root_node(p.text.strip(), idx)
            continue
        lvl = _heading_level(p)
        if lvl is not None:
            node = Node(title=p.text.strip(), level=lvl, start_idx=idx)
            while stack and stack[-1].level >= lvl:
                stack.pop()
            if stack: stack[-1].children.append(node)
            else: roots.append(node)
            stack.append(node)
        else:
            if stack:
                stack[-1].content_idxs.append(idx)
                if has_sectPr(p):   # stop at section end
                    stack.clear()
    return roots

def iter_nodes(nodes: List[Node]) -> List[Node]:
    out=[]
    def walk(n: Node):
        out.append(n)
        for c in n.children: walk(c)
    for r in nodes: walk(r)
    return out

# ----------------- safe ops -----------------
def clear_paragraph_text(paragraph: Paragraph):
    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)
    paragraph.text = ""

def remove_paragraph(paragraph: Paragraph):
    paragraph._element.getparent().remove(paragraph._element)

def remove_table(table: Table):
    table._element.getparent().remove(table._element)

def is_p(elm) -> bool:
    return elm.tag == qn("w:p")

def is_tbl(elm) -> bool:
    return elm.tag == qn("w:tbl")

def para_from_elm(doc: Document, elm) -> Paragraph:
    return Paragraph(elm, doc._body)

def table_from_elm(doc: Document, elm) -> Table:
    return Table(elm, doc._body)

# ----------------- body-style inference -----------------
def infer_body_style(doc: Document, node: Node):
    for i in node.content_idxs:
        if 0 <= i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            if _heading_level(p) is None and not has_sectPr(p) and p.text.strip():
                try: return p.style
                except Exception: pass
    try: return doc.styles["Normal"]
    except Exception: return None

def insert_paragraph_after(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try: new_para.style = style
        except Exception: pass
    if text: new_para.add_run(text)
    return new_para

# ----------------- core replace (BLOCK-AWARE) -----------------
def apply_content_to_node(doc: Document, node: Node, new_text: str, body_style=None):
    """
    Replace *direct* body under the node's heading (before first subheading or section break).
    Removes paragraphs and tables until a heading (any level) or sectPr paragraph.
    Preserves sectPr so headers/footers & pagination remain.
    """
    heading_p = doc.paragraphs[node.start_idx]
    cur = heading_p._p.getnext()

    # 1) Remove blocks until boundary
    while cur is not None:
        if is_p(cur):
            p = para_from_elm(doc, cur)
            if has_sectPr(p):
                clear_paragraph_text(p)  # keep section boundary
                break
            if _heading_level(p) is not None:
                break  # stop at first subheading/next heading
            nxt = cur.getnext()
            remove_paragraph(p)
            cur = nxt
            continue
        elif is_tbl(cur):
            nxt = cur.getnext()
            remove_table(table_from_elm(doc, cur))
            cur = nxt
            continue
        else:
            cur = cur.getnext()

    # 2) Insert new body paragraphs right after the heading
    blocks = [t.strip() for t in re.split(r"\n\s*\n", new_text or "") if t.strip()]
    cursor = heading_p
    for b in blocks:
        cursor = insert_paragraph_after(cursor, b, style=body_style)

# ----------------- driver -----------------
def apply_replacements(docx_in: str, json_in: str, docx_out: str, edit_front_matter: bool = False, debug: bool = False):
    doc = Document(docx_in)
    tree = build_tree(doc)
    all_nodes = iter_nodes(tree)

    title_to_nodes: Dict[str, List[Node]] = {}
    for n in all_nodes:
        title_to_nodes.setdefault(normalize_text(n.title), []).append(n)

    with open(json_in, "r", encoding="utf-8") as f:
        spec: Dict[str, Any] = json.load(f)

    replacements: List[Tuple[Node, str]] = []

    # top-level keys like {"Chapter X": {"__content__": "..."}}
    for k, v in spec.items():
        if isinstance(v, dict) and "__content__" in v:
            nodes = title_to_nodes.get(normalize_text(k), [])
            if not nodes:
                want = normalize_text(k)
                partial = [n for t, ns in title_to_nodes.items() for n in ns if want and want in t]
                if len(partial) == 1: nodes = partial
            for node in nodes:
                replacements.append((node, v["__content__"]))

    def collect(d: Dict[str, Any]):
        for k, v in d.items():
            if k == "__content__": continue
            nodes = title_to_nodes.get(normalize_text(k), [])
            if not nodes:
                want = normalize_text(k)
                partial = [n for t, ns in title_to_nodes.items() for n in ns if want and want in t]
                if len(partial) == 1: nodes = partial
            if not nodes: 
                if debug: print(f"[skip] key not matched: {k}")
                continue
            if isinstance(v, dict):
                if "__content__" in v:
                    for node in nodes: replacements.append((node, v["__content__"]))
                collect(v)
            elif isinstance(v, str):
                for node in nodes: replacements.append((node, v))
    collect(spec)

    # de-dupe & apply bottom-up
    seen=set(); uniq=[]
    for node, text in replacements:
        key=(id(node), text)
        if key in seen: continue
        seen.add(key); uniq.append((node, text))
    uniq.sort(key=lambda x: x[0].start_idx, reverse=True)

    if debug:
        print("Will apply to:")
        for node, _ in uniq:
            print(f" - {node.title} (H{node.level}) at {node.start_idx}")

    for node, text in uniq:
        # skip front-matter by default
        if is_front_matter_title(doc.paragraphs[node.start_idx]) and not edit_front_matter:
            if debug: print(f"[skip front-matter] {node.title}")
            continue
        style = infer_body_style(doc, node)
        apply_content_to_node(doc, node, text, body_style=style)

    doc.save(docx_out)

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Apply JSON replacements (block-aware; preserves headers/footers; inherits styles)")
    ap.add_argument("input_docx")
    ap.add_argument("input_json")
    ap.add_argument("output_docx")
    ap.add_argument("--edit-front-matter", action="store_true")
    ap.add_argument("--debug", action="store_true")
    args = ap.parse_args()
    apply_replacements(args.input_docx, args.input_json, args.output_docx,
                       edit_front_matter=args.edit_front_matter, debug=args.debug)
    print(f"Saved updated document to {args.output_docx}")
